"""Сборка v2 отчёта с настоящими Word-формулами (OMML)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from copy import deepcopy
import os, re

doc = Document()
for style in doc.styles:
    if style.name == 'Normal':
        f = style.font
        f.name = 'Times New Roman'
        f.size = Pt(13)

NS_M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
nsmap['m'] = NS_M

def m(tag): return f'{{{NS_M}}}{tag}'

def make_omml_run(text, italic=False, bold=False):
    """Создать <m:r><m:t>text</m:t></m:r>."""
    r = OxmlElement('m:r')
    if italic or bold:
        rpr = OxmlElement('m:rPr')
        # OMML стиль
        sty = OxmlElement('m:sty')
        if italic and bold: sty.set(qn('m:val'), 'bi')
        elif italic:        sty.set(qn('m:val'), 'i')
        elif bold:          sty.set(qn('m:val'), 'b')
        rpr.append(sty)
        r.append(rpr)
    t = OxmlElement('m:t')
    t.set(qn('xml:space'),'preserve')
    t.text = text
    r.append(t)
    return r

def make_frac(num_elements, den_elements):
    """Создать дробь <m:f>."""
    f = OxmlElement('m:f')
    fpr = OxmlElement('m:fPr')
    f.append(fpr)
    num = OxmlElement('m:num')
    den = OxmlElement('m:den')
    for e in num_elements: num.append(e)
    for e in den_elements: den.append(e)
    f.append(num); f.append(den)
    return f

def make_sup(base_elements, sup_elements):
    """e^x — <m:sSup>."""
    s = OxmlElement('m:sSup')
    e = OxmlElement('m:e'); sup = OxmlElement('m:sup')
    for x in base_elements: e.append(x)
    for x in sup_elements:  sup.append(x)
    s.append(e); s.append(sup)
    return s

def make_sub(base_elements, sub_elements):
    s = OxmlElement('m:sSub')
    e = OxmlElement('m:e'); sub = OxmlElement('m:sub')
    for x in base_elements: e.append(x)
    for x in sub_elements:  sub.append(x)
    s.append(e); s.append(sub)
    return s

def parse_formula(s):
    """Очень простой парсер: распознаёт _x (sub), ^x (sup), /текст/ как дробь Frac(num/den).
    Использует синтаксис: \\frac{num}{den}, x_{i}, x^{n}.
    Для упрощения — фиксированные шаблоны."""
    pass

def omml_paragraph(elements_list, align='left'):
    """Вставить параграф с OMML формулой."""
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'just':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    omath_para = OxmlElement('m:oMathPara')
    omath = OxmlElement('m:oMath')
    for e in elements_list:
        omath.append(e)
    omath_para.append(omath)
    p._element.append(omath_para)
    return p

def make_text(s, italic=False):
    """Просто текст в OMML (italic для переменных)."""
    return make_omml_run(s, italic=italic)

def make_var(name):
    return make_omml_run(name, italic=True)

def make_op(symbol):
    return make_omml_run(symbol, italic=False)

# ---- Удобные построители ----
def F_frac(num_str, den_str):
    """Дробь num/den, где числитель и знаменатель — простые строки."""
    return make_frac([make_text(num_str)], [make_text(den_str)])

def F_frac_complex(num_elems, den_elems):
    return make_frac(num_elems, den_elems)

def add_para(text, bold=False, italic=False, align=None, size=13):
    p = doc.add_paragraph()
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'just': p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size)
    r.bold=bold; r.italic=italic
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name='Times New Roman'
        run.font.color.rgb = RGBColor(0,0,0)

def add_image(path, width_cm=15.5):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(path, width=Cm(width_cm))

# ----- helpers для типовых формул -----
def add_eq_inline(par, omml_elements):
    """Встроить OMML в существующий параграф."""
    omath = OxmlElement('m:oMath')
    for e in omml_elements: omath.append(e)
    par._element.append(omath)

def add_eq_block(omml_elements, align='center'):
    p = doc.add_paragraph()
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omath_para = OxmlElement('m:oMathPara')
    omath = OxmlElement('m:oMath')
    for e in omml_elements: omath.append(e)
    omath_para.append(omath)
    p._element.append(omath_para)
    return p

# ============ Сборка отчёта ============

# Титульный лист
add_para('Министерство науки и высшего образования Российской Федерации',
         align='center', size=12)
add_para('Санкт-Петербургский государственный университет аэрокосмического приборостроения (ГУАП)',
         align='center', size=12, bold=True)
for _ in range(6): doc.add_paragraph()
add_para('ОТЧЁТ', align='center', size=20, bold=True)
add_para('по соревнованию по исследованию нелинейных систем автоматического управления',
         align='center', size=14)
add_para('20–22 мая 2026 г.', align='center', size=13)
for _ in range(8): doc.add_paragraph()
add_para('Участник: Самарин Дмитрий Васильевич', size=13)
add_para('Группа: 3510М', size=13)
add_para('Команда состоит из одного человека.', size=13)
add_para('Все 8 заданий выполнены лично Самариным Д.В.', size=13, italic=True)
for _ in range(6): doc.add_paragraph()
add_para('Санкт-Петербург, 2026', align='center', size=12)
doc.add_page_break()

add_heading('Распределение работ по заданиям', level=1)
add_para('Состав команды: 1 человек — Самарин Дмитрий Васильевич (гр. 3510М).', align='just')
table = doc.add_table(rows=9, cols=2)
table.style = 'Table Grid'
table.cell(0,0).text = '№ задания'
table.cell(0,1).text = 'Исполнитель'
for i in range(1,9):
    table.cell(i,0).text = f'Задание {i}'
    table.cell(i,1).text = 'Самарин Дмитрий Васильевич'
doc.add_page_break()

# =========================================================
# ЗАДАНИЕ 1
# =========================================================
add_heading('Задание 1. Преобразование структурной схемы в эквивалентную ПФ', level=1)
add_para('Исходная структурная схема (рис. 1) содержит звенья:', align='just')

# W1(s) = 0.7/(1.5s+1)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_eq_inline(p, [
    make_sub([make_var('W')], [make_text('1')]),
    make_text('(s) = '),
    make_frac([make_text('0,7')], [make_text('1,5s + 1')]),
    make_text(',  '),
    make_sub([make_var('W')], [make_text('2')]),
    make_text('(s) = '),
    make_frac([make_text('3030')], [make_text('0,01s')]),
    make_text(',  '),
    make_sub([make_var('W')], [make_text('3')]),
    make_text(' = 10,'),
])
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_eq_inline(p, [
    make_sub([make_var('W')], [make_text('4')]),
    make_text('(s) = '),
    make_frac([make_text('1')], [make_text('0,1s')]),
    make_text(',  '),
    make_var('H'),
    make_text('(s) = '),
    make_frac([make_text('1')], [make_text('0,05s + 1')]),
    make_text(',  K = 13.'),
])

add_heading('Метод 1. Эквивалентные структурные преобразования', level=2)
add_para('Шаг 1. Параллельное соединение W₁(s) с пропорциональным усилителем K = 13:',
         align='just')
add_eq_block([
    make_sub([make_var('W')], [make_text('пар')]),
    make_text('(s) = '),
    make_sub([make_var('W')], [make_text('1')]),
    make_text('(s) + 13 = '),
    make_frac([make_text('0,7')], [make_text('1,5s + 1')]),
    make_text(' + 13 = '),
    make_frac([make_text('19,5s + 13,7')], [make_text('1,5s + 1')]),
    make_text('.'),
])

add_para('Шаг 2. Внутренний контур с единичной отрицательной ОС:', align='just')
add_eq_block([
    make_sub([make_var('W')], [make_text('внутр')]),
    make_text('(s) = '),
    make_frac(
        [make_sub([make_var('W')],[make_text('2')]),
         make_sub([make_var('W')],[make_text('3')])],
        [make_text('1 + '),
         make_sub([make_var('W')],[make_text('2')]),
         make_sub([make_var('W')],[make_text('3')])]
    ),
    make_text(' = '),
    make_frac([make_text('30 300')], [make_text('0,01s + 30 300')]),
    make_text('.'),
])

add_para('Шаг 3. Последовательное соединение прямой ветви:', align='just')
add_eq_block([
    make_sub([make_var('W')], [make_text('прям')]),
    make_text('(s) = '),
    make_sub([make_var('W')], [make_text('пар')]),
    make_text('(s) · '),
    make_sub([make_var('W')], [make_text('внутр')]),
    make_text('(s) · '),
    make_sub([make_var('W')], [make_text('4')]),
    make_text('(s) = '),
    make_frac(
        [make_text('6 060 000 · (195s + 137)')],
        [make_text('s · (3s + 2) · (s + 3 030 000)')]
    ),
    make_text('.'),
])

add_para('Шаг 4. Замыкание внешней обратной связи через H(s):', align='just')
add_eq_block([
    make_sub([make_var('W')], [make_text('зам')]),
    make_text('(s) = '),
    make_frac(
        [make_sub([make_var('W')],[make_text('прям')]), make_text('(s)')],
        [make_text('1 + '),
         make_sub([make_var('W')],[make_text('прям')]),
         make_text('(s) · '), make_var('H'), make_text('(s)')]
    ),
    make_text('.'),
])

add_para('После приведения к общему знаменателю и упрощения получаем эквивалентную ПФ замкнутой системы:',
         align='just')
add_eq_block([
    make_sub([make_var('W')], [make_text('зам')]),
    make_text('(s) = '),
    make_frac(
        [make_text('1 181 700 000 · '),
         make_sup([make_var('s')],[make_text('2')]),
         make_text(' + 24 464 220 000 · s + 16 604 400 000')],
        [make_text('3 · '),
         make_sup([make_var('s')],[make_text('4')]),
         make_text(' + 9 090 062 · '),
         make_sup([make_var('s')],[make_text('3')]),
         make_text(' + 187 860 040 · '),
         make_sup([make_var('s')],[make_text('2')]),
         make_text(' + 23 755 200 000 · s + 16 604 400 000')]
    ),
    make_text('.'),
])
add_para('Статический коэффициент: W_зам(0) = 1. Система астатическая, отслеживает позиционные сигналы без статической ошибки.',
         align='just')

add_heading('Метод 2. Формула Мейсона', level=2)
add_para('Сигнальный граф содержит один путь P₁ и две петли:', align='just')
add_eq_block([
    make_sub([make_var('L')],[make_text('1')]),
    make_text(' = − '),
    make_sub([make_var('W')],[make_text('2')]),
    make_sub([make_var('W')],[make_text('3')]),
    make_text(',  '),
    make_sub([make_var('L')],[make_text('2')]),
    make_text(' = − '),
    make_sub([make_var('W')],[make_text('прям')]),
    make_text(' · '),
    make_var('H'),
    make_text(',  Δ = 1 − ('),
    make_sub([make_var('L')],[make_text('1')]),
    make_text(' + '),
    make_sub([make_var('L')],[make_text('2')]),
    make_text(') + '),
    make_sub([make_var('L')],[make_text('1')]),
    make_sub([make_var('L')],[make_text('2')]),
    make_text('.'),
])
add_para('Поскольку обе петли касаются прямого пути, Δ₁ = 1, и:', align='just')
add_eq_block([
    make_sub([make_var('W')], [make_text('зам')]),
    make_text(' = '),
    make_frac(
        [make_sub([make_var('P')],[make_text('1')]),
         make_text(' · '),
         make_sub([make_var('Δ')],[make_text('1')])],
        [make_var('Δ')]
    ),
    make_text('.'),
])
add_para('Численный результат идентичен методу 1 (контроль в скрипте task1_reduce.py).',
         align='just')

add_heading('Метод 3. Символьные вычисления в MATLAB', level=2)
add_para('Применяя функции feedback, series, parallel из Control System Toolbox:',
         align='just')
add_para('  W1 = 0.7/(1.5*s+1); W2 = 3030/(0.01*s); W3 = 10; W4 = 1/(0.1*s);',
         align='just', size=11)
add_para('  H = 1/(0.05*s+1); K13 = 13;', align='just', size=11)
add_para('  W_par = W1 + K13;', align='just', size=11)
add_para('  W_inner = feedback(W2*W3, 1);', align='just', size=11)
add_para('  W_forward = W_par * W_inner * W4;', align='just', size=11)
add_para('  W_cl = minreal(feedback(W_forward, H));', align='just', size=11)
add_para('Численный результат идентичен полученному вручную (см. task1_reduce.m).',
         align='just')
add_image('figures/task2_output.png')
add_para('Рис. 1.1. Переходная характеристика замкнутой системы при x = 100·1(t).',
         align='center', italic=True, size=11)

# =========================================================
# ЗАДАНИЕ 2
# =========================================================
doc.add_page_break()
add_heading('Задание 2. Установившаяся ошибка при x = 100·1(t)', level=1)
add_para('Аналитический расчёт. Изображение ошибки:', align='just', bold=True)
add_eq_block([
    make_var('E'), make_text('(s) = '),
    make_frac(
        [make_var('X'), make_text('(s)')],
        [make_text('1 + '),
         make_sub([make_var('W')],[make_text('разомк')]),
         make_text('(s)')]
    ),
    make_text(',  '), make_var('X'), make_text('(s) = '),
    make_frac([make_text('100')], [make_var('s')]),
    make_text('.'),
])
add_para('По теореме о конечном значении:', align='just')
add_eq_block([
    make_sub([make_var('e')], [make_text('уст')]),
    make_text(' = '),
    make_text('lim'),
    make_sub([make_text('')],[make_text('s → 0')]),
    make_text(' s · '),
    make_var('E'), make_text('(s) = '),
    make_frac(
        [make_text('100')],
        [make_text('1 + '),
         make_sub([make_var('W')],[make_text('разомк')]),
         make_text('(0)')]
    ),
    make_text('.'),
])
add_para('Петлевая ПФ:', align='just', bold=True)
add_eq_block([
    make_sub([make_var('W')], [make_text('разомк')]),
    make_text('(s) = '),
    make_frac(
        [make_text('23 634 000 000 · s + 16 604 400 000')],
        [make_text('3 · '),
         make_sup([make_var('s')],[make_text('4')]),
         make_text(' + 9 090 062 · '),
         make_sup([make_var('s')],[make_text('3')]),
         make_text(' + 187 860 040 · '),
         make_sup([make_var('s')],[make_text('2')]),
         make_text(' + 121 200 000 · s')]
    ),
    make_text('.'),
])
add_para('В знаменателе свободный множитель s — в петле есть интегратор. Следовательно W_разомк(s)|_{s→0} → ∞, и:',
         align='just')
add_eq_block([
    make_sub([make_var('e')],[make_text('уст')]),
    make_text(' = '),
    make_frac([make_text('100')],[make_text('1 + ∞')]),
    make_text(' = 0.'),
])
add_para('Моделирование. Симуляция в Simulink (task2_model.slx) и в Python (task2_error.py). График ошибки:',
         align='just', bold=True)
add_image('figures/task2_error.png')
add_para('Рис. 2.1. Переходный процесс ошибки e(t) при x(t) = 100·1(t).',
         align='center', italic=True, size=11)
add_para('Численное значение e_уст ≈ 3·10⁻¹¹ — машинный ноль, подтверждает аналитику.',
         align='just')

# =========================================================
# ЗАДАНИЕ 3
# =========================================================
doc.add_page_break()
add_heading('Задание 3. Синтез линейного регулятора', level=1)
add_para('Требования: y_уст = 80, σ ≤ 30 %, t_пп ≤ 7 с при x = 100·1(t). Объект W(s) — замкнутая ПФ из Task 1 (W(0) = 1).',
         align='just')

add_heading('Метод 1. П-регулятор по условию статики', level=2)
add_para('Замкнутая система:', align='just')
add_eq_block([
    make_frac(
        [make_var('y')],
        [make_var('x')]
    ),
    make_text(' = '),
    make_frac(
        [make_sub([make_var('K')],[make_text('p')]),
         make_text(' · '), make_var('W')],
        [make_text('1 + '),
         make_sub([make_var('K')],[make_text('p')]),
         make_text(' · '), make_var('W')]
    ),
    make_text('.'),
])
add_para('При W(0) = 1 условие y_уст = 80 при x = 100:', align='just')
add_eq_block([
    make_frac(
        [make_sub([make_var('K')],[make_text('p')])],
        [make_text('1 + '),
         make_sub([make_var('K')],[make_text('p')])]
    ),
    make_text(' = 0,8  ⇒  '),
    make_sub([make_var('K')],[make_text('p')]),
    make_text(' = 4.'),
])
add_para('Результат: y_уст = 80,00; σ = 20,78 %; t_пп = 0,07 с — все требования выполнены.',
         align='just')

add_heading('Метод 2. ПД-регулятор', level=2)
add_eq_block([
    make_var('C'), make_text('(s) = '),
    make_sub([make_var('K')],[make_text('p')]),
    make_text(' · (1 + '),
    make_sub([make_var('T')],[make_text('d')]),
    make_text(' · s),  '),
    make_sub([make_var('K')],[make_text('p')]),
    make_text(' = 4,  '),
    make_sub([make_var('T')],[make_text('d')]),
    make_text(' = 0,1.'),
])
add_para('Дифференциальная составляющая увеличивает запас по фазе. Результат: σ = 22,69 %; t_пп = 0,18 с.',
         align='just')

add_heading('Метод 3. Предкомпенсатор + П-регулятор с большим K_p', level=2)
add_eq_block([
    make_sub([make_var('K')],[make_text('p')]),
    make_text(' = 20,  '),
    make_var('prescale'),
    make_text(' = '),
    make_frac(
        [make_text('0,8 · (1 + '),
         make_sub([make_var('K')],[make_text('p')]),
         make_text(')')],
        [make_sub([make_var('K')],[make_text('p')])]
    ),
    make_text(' ≈ 0,84.'),
])
add_para('Результат: σ = 4,73 %; t_пп ≈ 0,003 с — минимальное перерегулирование, быстрейший процесс.',
         align='just')
add_image('figures/task3_compare.png')
add_para('Рис. 3.1. Переходные процессы для трёх методов синтеза.',
         align='center', italic=True, size=11)

# =========================================================
# ЗАДАНИЕ 4
# =========================================================
doc.add_page_break()
add_heading('Задание 4. Аппроксимация нелинейной характеристики F(e)', level=1)
add_para('Точки таблицы симметричны: F(−e) = F(e). Аппроксимируем F(|e|).',
         align='just')

add_heading('Метод 1. Кусочно-линейная интерполяция', level=2)
add_para('Реализация: блок 1-D Lookup Table в Simulink или функция interp1 в MATLAB.',
         align='just')

add_heading('Метод 2. Полиномиальная аппроксимация (3-я степень)', level=2)
add_eq_block([
    make_var('F'), make_text('(e) = 8,33·10'),
    make_sup([make_text('')],[make_text('−5')]),
    make_text(' · |e|'),
    make_sup([make_text('')],[make_text('3')]),
    make_text(' − 6,49·10'),
    make_sup([make_text('')],[make_text('−3')]),
    make_text(' · |e|'),
    make_sup([make_text('')],[make_text('2')]),
    make_text(' + 0,1722 · |e| − 0,020.'),
])

add_heading('Метод 3. Степенная аппроксимация', level=2)
add_eq_block([
    make_var('F'), make_text('(e) = 0,1864 · |e|'),
    make_sup([make_text('')],[make_text('0,6608')]),
    make_text('.'),
])
add_para('Получена линеаризацией в логарифмических координатах: log F = b·log|e| + log a.',
         align='just')
add_image('figures/task4_approx.png')
add_para('Рис. 4.1. Сопоставление аппроксимаций и табличных точек.',
         align='center', italic=True, size=11)
add_para('Влияние аппроксимации на переходный процесс:', align='just', bold=True)
add_image('figures/task4_compare.png')
add_para('Рис. 4.2. Переходные процессы при разных аппроксимациях.',
         align='center', italic=True, size=11)
add_para('Вывод: при |e| > 50 полиномиальная аппроксимация расходится (выходит за диапазон таблицы), что меняет характер процесса. Кусочно-линейная (с насыщением) и степенная — наиболее предсказуемые.',
         align='just')

# =========================================================
# ЗАДАНИЕ 5
# =========================================================
doc.add_page_break()
add_heading('Задание 5. Нелинейный логический регулятор', level=1)
add_para('Параметры: Δ = 1, k = 1. Регулятор содержит ветви:', align='just')

add_eq_block([
    make_sub([make_var('g')],[make_text('1')]),
    make_text('(e) = '),
])
# Conditional notation — render as case-like using \cases (in OMML — m:eqArr)
def make_cases(rows):
    """rows = [(lhs_elems, rhs_elems), ...]"""
    d = OxmlElement('m:d')
    dPr = OxmlElement('m:dPr')
    begChr = OxmlElement('m:begChr'); begChr.set(qn('m:val'), '{')
    endChr = OxmlElement('m:endChr'); endChr.set(qn('m:val'), '')
    sepChr = OxmlElement('m:sepChr'); sepChr.set(qn('m:val'), '')
    dPr.append(begChr); dPr.append(endChr); dPr.append(sepChr)
    d.append(dPr)
    e_el = OxmlElement('m:e')
    arr = OxmlElement('m:eqArr')
    for lhs, rhs in rows:
        sub_e = OxmlElement('m:e')
        for x in lhs: sub_e.append(x)
        for x in rhs: sub_e.append(x)
        arr.append(sub_e)
    e_el.append(arr)
    d.append(e_el)
    return d

add_eq_block([
    make_sub([make_var('g')],[make_text('1')]),
    make_text('(e) = '),
    make_cases([
        ([make_frac([make_text('2s + 1')],[make_text('5s + 1')]),
          make_text(' · e,   ')],
         [make_text('e > 10,')]),
        ([make_frac([make_text('2s + 1')],[make_text('3s + 1')]),
          make_text(' · e,   ')],
         [make_text('e < 10.')]),
    ])
])

add_eq_block([
    make_var('z'), make_text('(e) = '),
    make_cases([
        ([make_frac([make_text('18')],[make_text('0,05s + 1')]),
          make_text(' · '),
          make_frac([make_text('2s + 1')],[make_text('5s + 1')]),
          make_text(' · e,   ')],
         [make_text('|'),
          make_sub([make_var('g')],[make_text('1')]),
          make_text('| > 7,')]),
        ([make_text('5 · '),
          make_frac([make_text('2s + 1')],[make_text('3s + 1')]),
          make_text(' · e,   ')],
         [make_text('|'),
          make_sub([make_var('g')],[make_text('1')]),
          make_text('| ≤ 7.')]),
    ])
])

add_eq_block([
    make_var('F'), make_text('(e) = '),
    make_cases([
        ([make_text('0,   ')], [make_text('|z| ≤ Δ,')]),
        ([make_var('k'), make_text(' · z − Δ · sign(z),   ')],
         [make_text('|z| > Δ.')]),
    ])
])

add_para('Реализация в Simulink: динамические звенья — отдельные Transfer Fcn блоки '
         '(параллельно вычисляются), их выходы поступают в MATLAB Function-блок, '
         'который реализует логику переключений и зону нечувствительности.',
         align='just')
add_image('figures/task5_nlreg.png')
add_para('Рис. 5.1. Переходный процесс y(t) и управляющий сигнал F(e) при x = 100·1(t).',
         align='center', italic=True, size=11)
add_para('Полученные показатели: y_уст = 92,9; σ = 2,35 %; t_пп (по 5% трубке) ≈ 1,8 с.',
         align='just')

# =========================================================
# ЗАДАНИЕ 6
# =========================================================
doc.add_page_break()
add_heading('Задание 6. ПИД с ZOH в ОС и задержкой τ = 1,5 с', level=1)
add_para('Требования: σ ≤ 30 %, t_пп ≤ 10 с, y_уст = 20 при x = 100·1(t).',
         align='just')
add_para('Структура: x → масштаб 0,2 → сумматор → ПИД → задержка τ → объект W → y; обратная связь y через ZOH (T = 0,1 с) на сумматор.',
         align='just')
add_para('ПИД-регулятор:', align='just', bold=True)
add_eq_block([
    make_var('C'), make_text('(s) = '),
    make_sub([make_var('K')],[make_text('p')]),
    make_text(' + '),
    make_frac(
        [make_sub([make_var('K')],[make_text('i')])],
        [make_var('s')]
    ),
    make_text(' + '),
    make_sub([make_var('K')],[make_text('d')]),
    make_text(' · s.'),
])
add_para('Подобранные параметры:', align='just', bold=True)
add_eq_block([
    make_sub([make_var('K')],[make_text('p')]),
    make_text(' = 0,2;  '),
    make_sub([make_var('K')],[make_text('i')]),
    make_text(' = 0,35;  '),
    make_sub([make_var('K')],[make_text('d')]),
    make_text(' = 0.'),
])
add_para('Результат моделирования: y_уст = 20,00; σ = 0,00 %; t_пп = 6,65 с.',
         align='just')
add_image('figures/task6_pid.png')
add_para('Рис. 6.1. Переходный процесс системы с ПИД, задержкой и ZOH в ОС.',
         align='center', italic=True, size=11)

# =========================================================
# ЗАДАНИЕ 7
# =========================================================
doc.add_page_break()
add_heading('Задание 7. Реализация АИМ-2 через MATLAB Function', level=1)
add_para('АИМ-2 вырезает участки непрерывного сигнала на интервалах [nT, nT + γT), вне — 0. T = 1 с, γ = 0,6.',
         align='just')
add_para('Математическая модель:', align='just', bold=True)
add_eq_block([
    make_sup([make_var('x')],[make_text('*')]),
    make_text('(t) = '),
    make_cases([
        ([make_var('x'), make_text('(t),   ')],
         [make_text('mod(t, T) < γT,')]),
        ([make_text('0,   ')],
         [make_text('mod(t, T) ≥ γT.')]),
    ])
])
add_para('Тело MATLAB Function-блока:', align='just', bold=True)
add_para('function y = aim2(u, t)', align='just', size=11)
add_para('    T = 1.0;  gamma = 0.6;', align='just', size=11)
add_para('    phase = mod(t, T);', align='just', size=11)
add_para('    if phase < gamma*T', align='just', size=11)
add_para('        y = u;', align='just', size=11)
add_para('    else', align='just', size=11)
add_para('        y = 0;', align='just', size=11)
add_para('    end', align='just', size=11)
add_para('end', align='just', size=11)
add_para('Сигнал t берётся от блока Clock. Входной сигнал — синусоида 0,3 Гц.',
         align='just')
add_image('figures/task7_aim2.png')
add_para('Рис. 7.1. Сигнал на входе (пунктир) и выходе (сплошная линия) АИМ-2.',
         align='center', italic=True, size=11)

# =========================================================
# ЗАДАНИЕ 8
# =========================================================
doc.add_page_break()
add_heading('Задание 8. Система нелинейных ОДУ для ДПТ с регулируемым потоком', level=1)
add_para('Переменные состояния: I_я (ток якоря), I_в (ток обмотки возбуждения), ω_дв (угловая скорость), α (угол поворота).',
         align='just')

add_para('Якорная цепь:', align='just', bold=True)
add_eq_block([
    make_sub([make_var('L')],[make_text('я')]),
    make_text(' · '),
    make_frac(
        [make_text('d'),
         make_sub([make_var('I')],[make_text('я')])],
        [make_text('dt')]
    ),
    make_text(' + '),
    make_sub([make_var('R')],[make_text('я')]),
    make_text(' · '),
    make_sub([make_var('I')],[make_text('я')]),
    make_text(' = '),
    make_sub([make_var('U')],[make_text('я')]),
    make_text(' − '),
    make_sub([make_var('E')],[make_text('я')]),
    make_text(',   '),
    make_sub([make_var('E')],[make_text('я')]),
    make_text(' = C · Φ · '),
    make_sub([make_var('ω')],[make_text('дв')]),
    make_text('.'),
])

add_para('Цепь обмотки возбуждения:', align='just', bold=True)
add_eq_block([
    make_sub([make_var('L')],[make_text('в')]),
    make_text(' · '),
    make_frac(
        [make_text('d'),
         make_sub([make_var('I')],[make_text('в')])],
        [make_text('dt')]
    ),
    make_text(' + '),
    make_sub([make_var('R')],[make_text('в')]),
    make_text(' · '),
    make_sub([make_var('I')],[make_text('в')]),
    make_text(' = '),
    make_sub([make_var('U')],[make_text('в')]),
    make_text(',   Φ = '),
    make_sub([make_var('K')],[make_text('Φ')]),
    make_text(' · '),
    make_sub([make_var('I')],[make_text('в')]),
    make_text('.'),
])

add_para('Электромагнитный момент и уравнение баланса моментов:', align='just', bold=True)
add_eq_block([
    make_sub([make_var('M')],[make_text('эм')]),
    make_text(' = C · Φ · '),
    make_sub([make_var('I')],[make_text('я')]),
    make_text(',   '),
    make_var('J'),
    make_text(' · '),
    make_frac(
        [make_text('d'),
         make_sub([make_var('ω')],[make_text('дв')])],
        [make_text('dt')]
    ),
    make_text(' = '),
    make_sub([make_var('M')],[make_text('эм')]),
    make_text(' − '),
    make_frac(
        [make_sub([make_var('M')],[make_text('н')])],
        [make_var('i'), make_text(' · '), make_var('η')]
    ),
    make_text('.'),
])

add_para('Кинематика:', align='just', bold=True)
add_eq_block([
    make_var('ω'),
    make_text(' = '),
    make_frac(
        [make_sub([make_var('ω')],[make_text('дв')])],
        [make_var('i')]
    ),
    make_text(',   '),
    make_frac(
        [make_text('dα')],
        [make_text('dt')]
    ),
    make_text(' = '),
    make_var('ω'),
    make_text('.'),
])

add_para('Итоговая система четырёх нелинейных ОДУ:', align='just', bold=True)
add_eq_block([
    make_frac(
        [make_text('d'),
         make_sub([make_var('I')],[make_text('я')])],
        [make_text('dt')]
    ),
    make_text(' = '),
    make_frac([make_text('1')],
              [make_sub([make_var('L')],[make_text('я')])]),
    make_text(' · ('),
    make_sub([make_var('U')],[make_text('я')]),
    make_text(' − '),
    make_sub([make_var('R')],[make_text('я')]),
    make_text(' · '),
    make_sub([make_var('I')],[make_text('я')]),
    make_text(' − C · Φ · '),
    make_sub([make_var('ω')],[make_text('дв')]),
    make_text('),'),
])
add_eq_block([
    make_frac(
        [make_text('d'),
         make_sub([make_var('I')],[make_text('в')])],
        [make_text('dt')]
    ),
    make_text(' = '),
    make_frac([make_text('1')],
              [make_sub([make_var('L')],[make_text('в')])]),
    make_text(' · ('),
    make_sub([make_var('U')],[make_text('в')]),
    make_text(' − '),
    make_sub([make_var('R')],[make_text('в')]),
    make_text(' · '),
    make_sub([make_var('I')],[make_text('в')]),
    make_text('),'),
])
add_eq_block([
    make_frac(
        [make_text('d'),
         make_sub([make_var('ω')],[make_text('дв')])],
        [make_text('dt')]
    ),
    make_text(' = '),
    make_frac([make_text('1')], [make_var('J')]),
    make_text(' · (C · Φ · '),
    make_sub([make_var('I')],[make_text('я')]),
    make_text(' − '),
    make_frac(
        [make_sub([make_var('M')],[make_text('н')])],
        [make_var('i'), make_text(' · '), make_var('η')]
    ),
    make_text('),'),
])
add_eq_block([
    make_frac(
        [make_text('dα')],
        [make_text('dt')]
    ),
    make_text(' = '),
    make_frac(
        [make_sub([make_var('ω')],[make_text('дв')])],
        [make_var('i')]
    ),
    make_text('.'),
])

add_para('Нелинейность возникает из произведений Φ · ω_дв и Φ · I_я, поскольку Φ изменяется через I_в.',
         align='just')

# Заключение
doc.add_page_break()
add_heading('Заключение', level=1)
add_para('Все 8 заданий выполнены:', align='just')
add_para('1. Структурная схема рис. 1 свёрнута в эквивалентную ПФ тремя методами '
         '(эквивалентные преобразования, формула Мейсона, символьные вычисления в MATLAB).',
         align='just')
add_para('2. Аналитически и моделированием установлено e_уст = 0 (астатизм 1-го порядка).',
         align='just')
add_para('3. Синтезированы три линейных регулятора, все удовлетворяют σ ≤ 30 %, t_пп ≤ 7 с, y_уст = 80.',
         align='just')
add_para('4. Выполнены три аппроксимации F(e), показано их различное влияние на переходный процесс.',
         align='just')
add_para('5. Нелинейный логический регулятор: t_пп при 5 % трубке ≈ 1,8 с.', align='just')
add_para('6. ПИД-регулятор (Kp = 0,2, Ki = 0,35) с задержкой 1,5 с и ZOH: σ ≈ 0 %, t_пп ≈ 6,65 с.',
         align='just')
add_para('7. Реализована модель АИМ-2 через MATLAB Function в Simulink.', align='just')
add_para('8. Записана система нелинейных ОДУ для ДПТ с регулируемым потоком возбуждения.',
         align='just')
add_para('Все Simulink-модели и m-скрипты приложены к отчёту (папки models/, scripts/).',
         align='just')

out = 'Отчёт_Самарин_ГУАП_3510М_v2.docx'
doc.save(out)
print('Сохранено:', out)
