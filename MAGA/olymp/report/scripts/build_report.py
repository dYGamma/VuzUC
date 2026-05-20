"""Сборка финального отчёта Word (Самарин Д.В., ГУАП, 3510М)."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# Стиль документа
for style in doc.styles:
    if style.name == 'Normal':
        f = style.font
        f.name = 'Times New Roman'
        f.size = Pt(13)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_para(text, bold=False, italic=False, align=None, size=13):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'just':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def add_image(path, width_cm=15.5):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(path, width=Cm(width_cm))

# =================== Титульный лист ===================
add_para('Министерство науки и высшего образования Российской Федерации',
         align='center', size=12)
add_para('Санкт-Петербургский государственный университет аэрокосмического приборостроения (ГУАП)',
         align='center', size=12, bold=True)
for _ in range(6): doc.add_paragraph()
add_para('ОТЧЁТ', align='center', size=20, bold=True)
add_para('по соревнованию по исследованию нелинейных систем автоматического управления',
         align='center', size=14)
add_para('20–22 мая 2026 г.', align='center', size=13)
for _ in range(8): doc.add_paragraph()
add_para('Участник: Самарин Дмитрий Васильевич', size=13)
add_para('Группа: 3510М', size=13)
add_para('Команда состоит из одного человека.', size=13)
add_para('Все 8 заданий выполнены лично Самариным Д.В.', size=13, italic=True)
for _ in range(6): doc.add_paragraph()
add_para('Санкт-Петербург, 2026', align='center', size=12)

doc.add_page_break()

# =================== Распределение работ ===================
add_heading('Распределение работ по заданиям', level=1)
add_para('Состав команды: 1 человек — Самарин Дмитрий Васильевич (гр. 3510М).',
         align='just')
add_para('Распределение заданий между участниками:', align='just', bold=True)
table = doc.add_table(rows=9, cols=2)
table.style = 'Table Grid'
table.cell(0,0).text = '№ задания'
table.cell(0,1).text = 'Исполнитель'
for i in range(1, 9):
    table.cell(i,0).text = f'Задание {i}'
    table.cell(i,1).text = 'Самарин Дмитрий Васильевич'

doc.add_page_break()

# =================== ЗАДАНИЕ 1 ===================
add_heading('Задание 1. Преобразование структурной схемы в эквивалентную ПФ', level=1)

add_para('Исходная структурная схема (рис.1) содержит следующие звенья:', align='just')
add_para('  W₁(s) = 0,7 / (1,5s + 1) — корректирующий фильтр на входе;', align='just')
add_para('  W₂(s) = 3030 / (0,01s) — интегральный регулятор;', align='just')
add_para('  W₃ = 10 — усилитель;', align='just')
add_para('  W₄(s) = 1 / (0,1s) — механический интегратор;', align='just')
add_para('  H(s) = 1 / (0,05s + 1) — звено в обратной связи (датчик);', align='just')
add_para('  K = 13 — параллельная пропорциональная ветвь.', align='just')

add_heading('Метод 1. Эквивалентные преобразования структурной схемы', level=2)

add_para('Шаг 1. Параллельное соединение звена W₁(s) с пропорциональным '
         'усилителем K = 13 (оба сигнала поступают на сумматор «++»):',
         align='just')
add_para('  W_пар(s) = W₁(s) + 13 = 0,7/(1,5s+1) + 13 = (19,5s + 13,7) / (1,5s + 1).',
         align='center', italic=True)

add_para('Шаг 2. Внутренний контур с единичной отрицательной обратной связью '
         'охватывает последовательное соединение W₂(s)·W₃:',
         align='just')
add_para('  W_внутр(s) = W₂·W₃ / (1 + W₂·W₃) = (30300/0,01s) / (1 + 30300/0,01s) = '
         '30300 / (0,01s + 30300).',
         align='center', italic=True)

add_para('Шаг 3. Последовательное соединение W_пар, W_внутр и W₄:',
         align='just')
add_para('  W_прям(s) = W_пар(s) · W_внутр(s) · W₄(s) = '
         '6 060 000·(195s + 137) / (s·(3s+2)·(s + 3 030 000)).',
         align='center', italic=True)

add_para('Шаг 4. Замыкание внешней обратной связи через H(s):',
         align='just')
add_para('  W_зам(s) = W_прям(s) / (1 + W_прям(s)·H(s)).',
         align='center', italic=True)

add_para('После приведения к общему знаменателю и упрощения получаем эквивалентную '
         'передаточную функцию замкнутой системы:',
         align='just')
add_para('  W_зам(s) = (1 181 700 000·s² + 24 464 220 000·s + 16 604 400 000) /',
         align='center', italic=True)
add_para('  (3·s⁴ + 9 090 062·s³ + 187 860 040·s² + 23 755 200 000·s + 16 604 400 000).',
         align='center', italic=True)
add_para('Статический коэффициент усиления: W_зам(0) = 1,0 — система астатическая, '
         'отслеживает позиционные сигналы без статической ошибки.',
         align='just')

add_heading('Метод 2. Формула Мейсона', level=2)
add_para('Сигнальный граф системы содержит:',  align='just')
add_para('  — один путь от входа x к выходу y: P₁ = W_пар · W₂·W₃/(1+W₂·W₃) · W₄;',
         align='just')
add_para('  — три петли: L₁ = −W₂·W₃ (внутренняя единичная), '
         'L₂ = −W_прям·H (внешняя через датчик), '
         'L₃ = 0 (петли не пересекаются).',
         align='just')
add_para('Определитель графа Δ = 1 − L₁ − L₂ + (L₁·L₂).',
         align='just')
add_para('W_зам = P₁·Δ₁ / Δ, где Δ₁ = 1 (петли касаются прямого пути).',
         align='just')
add_para('После подстановки и упрощения результат совпадает с методом 1 '
         '(см. контрольную проверку в скрипте task1_reduce.py).',
         align='just')

add_heading('Метод 3. Символьные вычисления в MATLAB', level=2)
add_para('Применение функций feedback, series, parallel из Control System Toolbox:',
         align='just')
add_para('  W1 = 0.7/(1.5*s+1); W2 = 3030/(0.01*s); W3 = 10;', align='just', size=11)
add_para('  W4 = 1/(0.1*s); H = 1/(0.05*s+1); K13 = 13;', align='just', size=11)
add_para('  W_par = W1 + K13;', align='just', size=11)
add_para('  W_inner = feedback(W2*W3, 1);', align='just', size=11)
add_para('  W_forward = W_par * W_inner * W4;', align='just', size=11)
add_para('  W_cl = minreal(feedback(W_forward, H));', align='just', size=11)
add_para('Результат идентичен полученному вручную (см. task1_reduce.m).',
         align='just')

add_para('Графически переходная характеристика W_зам(s) показана на рис. 1.1.',
         align='just')
add_image('figures/task2_output.png')
add_para('Рис. 1.1. Переходная характеристика замкнутой системы W_зам(s) при x = 100·1(t).',
         align='center', italic=True, size=11)

# =================== ЗАДАНИЕ 2 ===================
doc.add_page_break()
add_heading('Задание 2. Установившаяся ошибка при x = 100·1(t)', level=1)

add_para('Аналитический расчёт:', align='just', bold=True)
add_para('Образ ошибки E(s) = X(s) / (1 + W_разомк(s)), где X(s) = 100/s — '
         'изображение единичного скачка амплитуды 100.',
         align='just')
add_para('По теореме о конечном значении:',  align='just')
add_para('  e_уст = lim_{s→0} s·E(s) = 100 / (1 + W_разомк(0)).',
         align='center', italic=True)
add_para('Петлевая передаточная функция:',  align='just')
add_para('  W_разомк(s) = W_прям(s)·H(s) = (23 634 000 000·s + 16 604 400 000) /',
         align='center', italic=True)
add_para('  (3·s⁴ + 9 090 062·s³ + 187 860 040·s² + 121 200 000·s + 0).',
         align='center', italic=True)
add_para('Знаменатель имеет нулевой свободный член (множитель s) — в петле имеется '
         'свободный интегратор. Следовательно W_разомк(s)|_{s→0} → ∞, и:',
         align='just')
add_para('  e_уст = 100 / (1 + ∞) = 0.',
         align='center', italic=True, bold=True)

add_para('Моделирование:',  align='just', bold=True)
add_para('Симуляция выполнена в Simulink (модель models/task2_error.slx) и продублирована '
         'в Python (scripts/task2_error.py). График ошибки e(t) приведён ниже.',
         align='just')
add_image('figures/task2_error.png')
add_para('Рис. 2.1. Переходный процесс ошибки e(t) при x(t) = 100·1(t).',
         align='center', italic=True, size=11)
add_para('Численное значение e_уст в установе ≈ 3·10⁻¹¹ — машинный ноль, что '
         'подтверждает аналитический результат.',
         align='just')

# =================== ЗАДАНИЕ 3 ===================
doc.add_page_break()
add_heading('Задание 3. Синтез линейного регулятора', level=1)
add_para('Требования: y_уст = 80, перерегулирование σ ≤ 30 %, '
         'время переходного процесса t_пп ≤ 7 с при x = 100·1(t).',
         align='just')
add_para('Объект W(s) — замкнутая передаточная функция, полученная в задании 1, '
         'W(0) = 1 (астатизм по позиции).', align='just')

add_heading('Метод 1. П-регулятор (по условию статики)', level=2)
add_para('Замкнутая структура: y/x = K_p·W/(1 + K_p·W).',  align='just')
add_para('При W(0)=1: y_уст/x = K_p/(1+K_p) = 0,8 ⇒ K_p = 4.',  align='just')
add_para('Результат моделирования: y_уст = 80,00; σ = 20,78 %; t_пп = 0,07 с. '
         'Все требования выполнены с большим запасом.',
         align='just')

add_heading('Метод 2. ПД-регулятор', level=2)
add_para('Регулятор C(s) = K_p·(1 + T_d·s), K_p = 4, T_d = 0,1.',
         align='just')
add_para('Дифференциальная составляющая увеличивает запас по фазе. '
         'Результат: σ = 22,69 %; t_пп = 0,18 с.', align='just')

add_heading('Метод 3. Предкомпенсатор + П-регулятор с большим K_p', level=2)
add_para('Контур замыкается с большим K_p = 20 (для уменьшения времени), '
         'на входе ставится масштабирующий множитель prescale = 0,8·(1+K_p)/K_p ≈ 0,84, '
         'чтобы y_уст = 80.', align='just')
add_para('Результат: σ = 4,73 %; t_пп ≈ 0,003 с — минимальное перерегулирование, '
         'быстрейший процесс.', align='just')

add_image('figures/task3_compare.png')
add_para('Рис. 3.1. Переходные процессы для трёх методов синтеза. '
         'Все методы удовлетворяют требованиям задания.',
         align='center', italic=True, size=11)

# =================== ЗАДАНИЕ 4 ===================
doc.add_page_break()
add_heading('Задание 4. Аппроксимация нелинейной характеристики F(e)', level=1)
add_para('Точки таблицы 1 симметричны относительно начала координат — '
         'характеристика чётная по |e|. Это позволяет аппроксимировать F(|e|).',
         align='just')

add_heading('Метод 1. Кусочно-линейная интерполяция (Lookup Table)', level=2)
add_para('Реализация: блок «1-D Lookup Table» в Simulink либо функция interp1 в MATLAB. '
         'Самая точная аппроксимация по табличным точкам.',  align='just')

add_heading('Метод 2. Полиномиальная аппроксимация (3-я степень)', level=2)
add_para('F(e) = 8,33·10⁻⁵·|e|³ − 6,49·10⁻³·|e|² + 1,72·10⁻¹·|e| − 0,020. '
         'Полином сглаживает точки; на краях возможны выбросы.',  align='just')

add_heading('Метод 3. Степенная аппроксимация', level=2)
add_para('F(e) = 0,1864·|e|^0,6608. '
         'Получена линеаризацией в логарифмических координатах: log F = b·log|e| + log a.',
         align='just')

add_image('figures/task4_approx.png')
add_para('Рис. 4.1. Сопоставление трёх аппроксимаций и табличных точек.',
         align='center', italic=True, size=11)

add_para('Сравнение переходных процессов системы (рис. 2 ТЗ) при '
         'использовании каждой аппроксимации:',  align='just', bold=True)
add_image('figures/task4_compare.png')
add_para('Рис. 4.2. Переходные процессы при разных аппроксимациях F(e). '
         'Кусочно-линейная и степенная дают близкое поведение, '
         'полиномиальная — существенно отличающееся из-за поведения вне диапазона данных.',
         align='center', italic=True, size=11)

add_para('Вывод: при больших значениях ошибки (|e|>50) полиномиальная аппроксимация '
         'даёт расходящееся значение F(e), что меняет характер переходного процесса. '
         'Кусочно-линейная (с насыщением вне таблицы) и степенная — наиболее предсказуемые.',
         align='just')

# =================== ЗАДАНИЕ 5 ===================
doc.add_page_break()
add_heading('Задание 5. Синтез нелинейного логического регулятора', level=1)
add_para('Регулятор задан системой условных уравнений с динамическими звеньями. '
         'Параметры: Δ = 1, k = 1.',
         align='just')

add_para('Условные выражения внутри регулятора:', align='just', bold=True)
add_para('  — при e > 10: g₁(e) = (2s+1)/(5s+1)·e;', align='just')
add_para('  — при e < 10: g₁(e) = (2s+1)/(3s+1)·e;', align='just')
add_para('  — если |g₁| > 7: z = 18/(0,05s+1)·(2s+1)/(5s+1)·e;', align='just')
add_para('  — иначе:        z = 5·(2s+1)/(3s+1)·e;', align='just')
add_para('  — если |z| ≤ Δ: F = 0;', align='just')
add_para('  — иначе: F = k·z − Δ·sign(z).', align='just')

add_para('Реализация в Simulink:',  align='just', bold=True)
add_para('Динамические звенья (2s+1)/(5s+1), (2s+1)/(3s+1), 18/(0,05s+1) реализованы '
         'параллельно как блоки Transfer Fcn. Их выходы поступают в MATLAB Function '
         'блок, реализующий логику переключений и зону нечувствительности. Такая '
         'структура устраняет необходимость менять состояние фильтров в runtime.',
         align='just')

add_image('figures/task5_nlreg.png')
add_para('Рис. 5.1. Переходный процесс и сигнал F(e) при x = 100·1(t).',
         align='center', italic=True, size=11)
add_para('Полученные показатели:',  align='just', bold=True)
add_para('  y_уст = 92,9; перерегулирование σ = 2,35 %; t_пп (по 5% трубке) ≈ 1,8 с.',
         align='just')

# =================== ЗАДАНИЕ 6 ===================
doc.add_page_break()
add_heading('Задание 6. ПИД-регулятор с ZOH в ОС и задержкой τ = 1,5 с', level=1)
add_para('Требования: σ ≤ 30 %, t_пп ≤ 10 с, y_уст = 20 при x = 100·1(t).',
         align='just')
add_para('Структура: x → [сумматор] → e → ПИД → задержка τ=1,5 с → объект W → y; '
         'обратная связь от y через ZOH (T_ZOH = 0,1 с) на сумматор.',
         align='just')

add_para('Для получения y_уст = 20 при x = 100 уставка масштабируется коэффициентом 0,2. '
         'Поскольку ПИД содержит интегральную составляющую, на установе e = 0 и '
         'y_уст = 0,2·x = 20.',
         align='just')

add_para('Подобранные параметры (по сетке Kp ∈ [0,2…0,6], Ki ∈ [0,1…0,35]):',
         align='just', bold=True)
add_para('  K_p = 0,2;  K_i = 0,35;  K_d = 0.', align='center', italic=True)
add_para('Результат моделирования:', align='just', bold=True)
add_para('  y_уст = 20,00;  σ = 0,00 %;  t_пп = 6,65 с.', align='just')
add_image('figures/task6_pid.png')
add_para('Рис. 6.1. Переходный процесс системы с ПИД, задержкой 1,5 с и ZOH в ОС.',
         align='center', italic=True, size=11)

# =================== ЗАДАНИЕ 7 ===================
doc.add_page_break()
add_heading('Задание 7. Реализация АИМ-2 через MATLAB Function', level=1)
add_para('АИМ-2 «вырезает» отдельные участки непрерывного сигнала: на интервалах '
         '[nT, nT + γT) выход равен входу, вне этих интервалов выход равен нулю. '
         'Параметры: T = 1 с, γ = 0,6 (60 %).',
         align='just')

add_para('Содержимое MATLAB Function-блока:',  align='just', bold=True)
add_para("function y = aim2(u, t)", align='just', size=11)
add_para("    T = 1.0; gamma = 0.6;", align='just', size=11)
add_para("    phase = mod(t, T);", align='just', size=11)
add_para("    if phase < gamma*T", align='just', size=11)
add_para("        y = u;", align='just', size=11)
add_para("    else", align='just', size=11)
add_para("        y = 0;", align='just', size=11)
add_para("    end", align='just', size=11)
add_para("end", align='just', size=11)

add_para('Сигнал времени t подаётся в блок MATLAB Function от блока Clock в Simulink. '
         'В качестве входного сигнала u использован синусоидальный сигнал с частотой 0,3 Гц.',
         align='just')
add_image('figures/task7_aim2.png')
add_para('Рис. 7.1. Сигнал на входе АИМ-2 (пунктир) и на его выходе (сплошная линия). '
         'Затенённые области соответствуют интервалам открытого ключа.',
         align='center', italic=True, size=11)

# =================== ЗАДАНИЕ 8 ===================
doc.add_page_break()
add_heading('Задание 8. Система нелинейных ОДУ для ДПТ с изменяемым потоком возбуждения', level=1)

add_para('По структурной схеме рис.4 (двигатель постоянного тока с регулируемым '
         'потоком возбуждения) переменные состояния:', align='just')
add_para('  Iя — ток якоря;  Iв — ток обмотки возбуждения;', align='just')
add_para('  ω — угловая скорость вала двигателя; α — угол поворота вала.', align='just')

add_para('Якорная цепь:', align='just', bold=True)
add_para('  L_я · dI_я/dt + R_я · I_я = U_я − E_я,', align='center', italic=True)
add_para('  где E_я = C · Φ · ω — противо-ЭДС.', align='just')

add_para('Цепь обмотки возбуждения:', align='just', bold=True)
add_para('  L_в · dI_в/dt + R_в · I_в = U_в,', align='center', italic=True)
add_para('  Φ = K_Φ · I_в  (в линейном приближении магнитной характеристики).', align='just')

add_para('Уравнение баланса электромагнитного и нагрузочного моментов:', align='just', bold=True)
add_para('  J · dω_дв/dt = M_эм − M_н/(i·η),', align='center', italic=True)
add_para('  где M_эм = C · Φ · I_я.', align='just')

add_para('Кинематика:', align='just', bold=True)
add_para('  ω_вых = ω_дв / i  (редуктор с передаточным числом i);', align='just')
add_para('  dα/dt = ω_вых.', align='just')

add_para('Итоговая система четырёх нелинейных ОДУ:', align='just', bold=True)
add_para('  dI_я/dt = (1/L_я)·(U_я − R_я·I_я − C·Φ·ω_дв);', align='center', italic=True)
add_para('  dI_в/dt = (1/L_в)·(U_в − R_в·I_в);', align='center', italic=True)
add_para('  dω_дв/dt = (1/J)·(C·Φ·I_я − M_н/(i·η));', align='center', italic=True)
add_para('  dα/dt = ω_дв / i.', align='center', italic=True)
add_para('Алгебраические связи: Φ = K_Φ·I_в;  E_я = C·Φ·ω_дв;  M_эм = C·Φ·I_я.',
         align='just')
add_para('Нелинейность системы возникает из произведений Φ·ω_дв и Φ·I_я, поскольку '
         'поток возбуждения Φ изменяется через ток I_в.',
         align='just')

# =================== Заключение ===================
doc.add_page_break()
add_heading('Заключение', level=1)
add_para('В рамках соревнования выполнены все 8 конкурсных заданий по исследованию '
         'нелинейных систем автоматического управления:', align='just')
add_para('  1. Структурная схема рис.1 преобразована в эквивалентную ПФ тремя методами '
         '(эквивалентные преобразования, формула Мейсона, MATLAB feedback).',
         align='just')
add_para('  2. Аналитически и моделированием подтверждено, что система астатическая 1-го '
         'порядка по позиционному входу и e_уст = 0.', align='just')
add_para('  3. Синтезированы три линейных регулятора (П, ПД, П+предкомпенсатор), '
         'все удовлетворяют σ ≤ 30 %, t_пп ≤ 7 с, y_уст = 80.', align='just')
add_para('  4. Выполнены три аппроксимации нелинейной характеристики F(e); показано '
         'различие влияния каждой аппроксимации на переходный процесс.',
         align='just')
add_para('  5. Реализован нелинейный логический регулятор; t_пп при 5% трубке ≈ 1,8 с.',
         align='just')
add_para('  6. Подобран ПИД-регулятор (Kp=0,2, Ki=0,35) с учётом задержки 1,5 с и ZOH в '
         'обратной связи; σ ≈ 0 %, t_пп ≈ 6,65 с, y_уст = 20.', align='just')
add_para('  7. Реализована модель АИМ-2 через MATLAB Function в Simulink.',
         align='just')
add_para('  8. Записана система нелинейных ОДУ для ДПТ с изменяемым потоком возбуждения.',
         align='just')
add_para('Все Simulink-модели и m-скрипты приложены к отчёту.', align='just')

# Сохраняем
out = 'Отчёт_Самарин_ГУАП_3510М.docx'
doc.save(out)
print(f'Сохранено: {out}')
