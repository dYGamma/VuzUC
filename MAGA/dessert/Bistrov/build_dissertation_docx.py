# -*- coding: utf-8 -*-
"""
Сборка магистерской диссертации (введение + глава 1) в .docx.
Оформление по ГОСТ Р 7.0.11-2011 и требованиям ГУАП:
- Times New Roman 14, межстрочный 1,5
- Поля: левое 30 мм, правое 15 мм, верхнее/нижнее 20 мм
- Абзацный отступ 1,25 см, выравнивание по ширине
- Заголовки H1: ПРОПИСНЫМИ, полужирный, по центру
- Заголовки H2: с прописной, полужирный, выравнивание влево с абзацного отступа
- Нумерация страниц снизу по центру, со 2-й (титул без номера)
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

THESIS_TOPIC = (
    "Метод автоматизированного синтеза программного кода "
    "промышленных контроллеров на основе графовой модели "
    "схемы автоматизации"
)
STUDENT = "Самарин Дмитрий Викторович"
GROUP = "3510М"
SPECIALTY = "27.04.04 «Управление в технических системах»"

ROOT = Path(r"C:\VuzUC\MAGA\dessert\Bistrov")
INPUT_FILES = [ROOT / "01_Introduction.md", ROOT / "02_Chapter_1.md"]
OUTPUT = ROOT / "Dissertation_draft.docx"

# =============================================================
# Базовые утилиты форматирования
# =============================================================

doc = Document()

for section in doc.sections:
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(14)
rpr = normal.element.get_or_add_rPr()
rfonts = OxmlElement('w:rFonts')
rfonts.set(qn('w:ascii'), 'Times New Roman')
rfonts.set(qn('w:hAnsi'), 'Times New Roman')
rfonts.set(qn('w:cs'), 'Times New Roman')
rfonts.set(qn('w:eastAsia'), 'Times New Roman')
rpr.append(rfonts)

pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.first_line_indent = Cm(1.25)


def _apply_run_font(run, size=14, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), 'Times New Roman')
    rfonts.set(qn('w:hAnsi'), 'Times New Roman')
    rfonts.set(qn('w:cs'), 'Times New Roman')
    rpr.append(rfonts)


def add_paragraph(text, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
                  indent=True, size=14, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.first_line_indent = Cm(1.25) if indent else Cm(0)
    if text:
        # обработка inline **bold**
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2])
                _apply_run_font(r, size=size, bold=True)
            else:
                r = p.add_run(part)
                _apply_run_font(r, size=size, bold=bold)
    return p


def add_h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(18)
    pf.space_after = Pt(18)
    pf.first_line_indent = Cm(0)
    pf.keep_with_next = True
    r = p.add_run(text.upper())
    _apply_run_font(r, size=14, bold=True)


def add_h2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(1.25)
    pf.keep_with_next = True
    r = p.add_run(text)
    _apply_run_font(r, size=14, bold=True)


def add_h3(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(1.25)
    pf.keep_with_next = True
    r = p.add_run(text)
    _apply_run_font(r, size=14, bold=True)


def add_bullet(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.left_indent = Cm(1.25)
    pf.first_line_indent = Cm(0)
    parts = re.split(r'(\*\*[^*]+\*\*)', '— ' + text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            _apply_run_font(r, bold=True)
        else:
            r = p.add_run(part)
            _apply_run_font(r)


def add_numbered(text, idx):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.left_indent = Cm(1.25)
    pf.first_line_indent = Cm(0)
    prefix = f'{idx}. '
    parts = re.split(r'(\*\*[^*]+\*\*)', prefix + text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            _apply_run_font(r, bold=True)
        else:
            r = p.add_run(part)
            _apply_run_font(r)


def page_break():
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


# =============================================================
# Markdown parser (минимальный, под наши .md)
# =============================================================

def parse_markdown(path: Path):
    text = path.read_text(encoding='utf-8')

    # удалить HTML-комментарии
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)

    lines = text.split('\n')
    num_pattern = re.compile(r'^(\d+)\.\s+(.*)$')

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # Заголовки
        if line.startswith('### '):
            add_h3(line[4:].strip())
            i += 1
            continue
        if line.startswith('## '):
            add_h2(line[3:].strip())
            i += 1
            continue
        if line.startswith('# '):
            add_h1(line[2:].strip())
            i += 1
            continue

        # Маркированные списки: — text  /  - text
        if line.startswith('— ') or line.startswith('- '):
            txt = line[2:].strip()
            # склеить продолжения (строки с продолжением одного буллета)
            while i + 1 < len(lines):
                nxt = lines[i + 1].rstrip()
                if nxt and not nxt.startswith(('— ', '- ', '# ', '## ',
                                               '### ')) \
                        and not num_pattern.match(nxt) \
                        and nxt[0].islower():
                    txt += ' ' + nxt.strip()
                    i += 1
                else:
                    break
            add_bullet(txt.rstrip(';.').rstrip())
            i += 1
            continue

        # Нумерованные списки
        m = num_pattern.match(line)
        if m:
            idx = int(m.group(1))
            txt = m.group(2).strip()
            while i + 1 < len(lines):
                nxt = lines[i + 1].rstrip()
                if nxt and not num_pattern.match(nxt) \
                        and not nxt.startswith(('— ', '- ', '# ', '## ',
                                                '### ')) \
                        and nxt[0].islower():
                    txt += ' ' + nxt.strip()
                    i += 1
                else:
                    break
            add_numbered(txt, idx)
            i += 1
            continue

        # Обычный абзац — склеить многострочный
        txt = line
        while i + 1 < len(lines):
            nxt = lines[i + 1].rstrip()
            if nxt and not nxt.startswith(('— ', '- ', '# ', '## ',
                                           '### ')) \
                    and not num_pattern.match(nxt):
                txt += ' ' + nxt.strip()
                i += 1
            else:
                break
        add_paragraph(txt.strip())
        i += 1


# =============================================================
# Титульный лист по ГОСТ Р 7.0.11-2011
# =============================================================

add_paragraph('МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False)
add_paragraph('РОССИЙСКОЙ ФЕДЕРАЦИИ',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False)
add_paragraph('Федеральное государственное автономное образовательное '
              'учреждение высшего образования',
              align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, before=6)
add_paragraph('«САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ '
              'АЭРОКОСМИЧЕСКОГО ПРИБОРОСТРОЕНИЯ»',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False)
add_paragraph('', indent=False, before=12)
add_paragraph('Институт киберфизических систем',
              align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
add_paragraph('Кафедра № 31',
              align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
add_paragraph('', indent=False, before=24)
add_paragraph('На правах рукописи',
              align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
add_paragraph('', indent=False, before=12)
add_paragraph(STUDENT, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
              indent=False, before=24)
add_paragraph('', indent=False, before=12)
add_paragraph(THESIS_TOPIC.upper(), align=WD_ALIGN_PARAGRAPH.CENTER,
              bold=True, indent=False, before=12)
add_paragraph('', indent=False, before=12)
add_paragraph(f'Направление подготовки: {SPECIALTY}',
              align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
add_paragraph('', indent=False, before=12)
add_paragraph('МАГИСТЕРСКАЯ ДИССЕРТАЦИЯ',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False,
              before=12)
add_paragraph('на соискание квалификации магистра',
              align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
add_paragraph('', indent=False, before=48)
add_paragraph(f'Студент группы № {GROUP}: ______________ / {STUDENT} /',
              align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
add_paragraph('', indent=False, before=12)
add_paragraph('Научный руководитель: ______________ / ____________________ /',
              align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
add_paragraph('', indent=False, before=12)
add_paragraph('«Допустить к защите»', align=WD_ALIGN_PARAGRAPH.LEFT,
              indent=False)
add_paragraph('Заведующий кафедрой № 31: ______________ / ___________________ /',
              align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
add_paragraph('«____» ____________ 2027 г.', align=WD_ALIGN_PARAGRAPH.LEFT,
              indent=False)
add_paragraph('', indent=False, before=48)
add_paragraph('Санкт-Петербург', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
add_paragraph('2027', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

page_break()

# =============================================================
# Содержание (заглушка — будет дополнено по мере написания работы)
# =============================================================

add_h1('Содержание')
toc = [
    'ВВЕДЕНИЕ',
    'ГЛАВА 1. АНАЛИЗ ЗАДАЧ РАЗРАБОТКИ ПРОГРАММНОГО ОБЕСПЕЧЕНИЯ '
    'АВТОМАТИЗИРОВАННЫХ СИСТЕМ УПРАВЛЕНИЯ И ОБОСНОВАНИЕ ГРАФОВОГО '
    'ПОДХОДА К ПРЕДСТАВЛЕНИЮ СХЕМ АВТОМАТИЗАЦИИ',
    '   1.1 Эволюция средств автоматизированного проектирования систем '
    'управления',
    '   1.2 Стандарт IEC 61131-3 и современные среды разработки '
    'программного обеспечения промышленных контроллеров',
    '   1.3 «Семантический разрыв» между средами проектирования аппаратной '
    'и программной частей АСУ ТП',
    '   1.4 Обзор существующих подходов к интеграции ECAD- и PLC-систем',
    '   1.5 Графовые модели в представлении инженерных систем',
    '   1.6 Модельно-ориентированное проектирование и автоматическая '
    'генерация кода',
    '   1.7 Нормативная база и контекст импортозамещения в Российской '
    'Федерации',
    '   1.8 Постановка научной задачи диссертационного исследования',
    '   1.9 Выводы по главе',
    'ГЛАВА 2. ГРАФОВАЯ МОДЕЛЬ СХЕМЫ АВТОМАТИЗАЦИИ И ТИПОВЫЕ ПАТТЕРНЫ '
    'ОБОРУДОВАНИЯ (в разработке)',
    'ГЛАВА 3. МЕТОД СИНТЕЗА ПРОГРАММНОГО КОДА ПЛК СТАНДАРТА IEC 61131-3 '
    '(в разработке)',
    'ГЛАВА 4. ЭКСПЕРИМЕНТАЛЬНОЕ ИССЛЕДОВАНИЕ И ОЦЕНКА ЭФФЕКТИВНОСТИ '
    'МЕТОДА (в разработке)',
    'ЗАКЛЮЧЕНИЕ (в разработке)',
    'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ (в разработке)',
    'ПРИЛОЖЕНИЯ (в разработке)',
]
for t in toc:
    add_paragraph(t, align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)

page_break()

# =============================================================
# Парсим введение + главу 1 из .md
# =============================================================

for idx, md_path in enumerate(INPUT_FILES):
    parse_markdown(md_path)
    if idx < len(INPUT_FILES) - 1:
        page_break()


# =============================================================
# Номера страниц со 2-й
# =============================================================

def add_page_numbers(doc):
    sect = doc.sections[0]
    sect.different_first_page_header_footer = True

    footer = sect.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    run = p.add_run()
    _apply_run_font(run, size=12)
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)

    sect.first_page_footer.paragraphs[0].text = ''


add_page_numbers(doc)

doc.save(str(OUTPUT))
print(f'Saved: {OUTPUT}')
